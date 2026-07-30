"""Plain-text extraction from resume files, for deterministic JD matching.

Hard rule: extraction must NEVER raise into the caller. A corrupt, truncated,
or otherwise unreadable file yields an empty string, so an upload can never be
turned into a 500 by a malformed document. No file content is ever logged (only
the kind and a generic failure note), because a resume is personal data.
"""

import io
import logging
import re

logger = logging.getLogger(__name__)

_WHITESPACE_RE = re.compile(r"\s+")
# A LaTeX comment: an unescaped `%` to end of line. `\%` is a literal percent.
_TEX_COMMENT_RE = re.compile(r"(?<!\\)%.*")
_TEX_HREF_RE = re.compile(r"\\href\s*\{[^{}]*\}\s*\{([^{}]*)\}")
_TEX_URL_RE = re.compile(r"\\url\s*\{([^{}]*)\}")
_TEX_ENV_RE = re.compile(r"\\(?:begin|end)\s*\{[^{}]*\}")
# A command carrying a braced argument: keep the argument, drop the command.
_TEX_CMD_ARG_RE = re.compile(r"\\[a-zA-Z@]+\s*(?:\[[^\]]*\])?\s*\{([^{}]*)\}")
# A bare command (\item, \hfill, \newpage, ...): drop it, optional bracket opts.
_TEX_CMD_BARE_RE = re.compile(r"\\[a-zA-Z@]+(?:\[[^\]]*\])?")
# An escaped special (\& \% \_ \# \$ \{ \}): keep the literal character.
_TEX_ESCAPED_RE = re.compile(r"\\([&%_#$}{])")


def normalize_whitespace(text: str) -> str:
    """Collapse every run of whitespace to a single space and trim the ends."""
    return _WHITESPACE_RE.sub(" ", text).strip()


def extract_text(kind: str, data: bytes) -> str:
    """Best-effort plain text for a resume of the given kind.

    Returns "" on any failure or unknown kind; never raises.
    """
    if not data:
        return ""
    try:
        if kind == "pdf":
            return _extract_pdf(data)
        if kind == "docx":
            return _extract_docx(data)
        if kind == "tex":
            return _extract_tex(data)
    except Exception:
        # The text is optional; a bad file must not break the request.
        logger.warning("resume text extraction failed for kind=%s", kind)
    return ""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            # One unreadable page must not lose the rest of the document.
            page_text = ""
        if page_text:
            parts.append(page_text)
    return normalize_whitespace("\n".join(parts))


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return normalize_whitespace("\n".join(parts))


def _extract_tex(data: bytes) -> str:
    """Strip LaTeX markup down to readable prose.

    This is deliberately a lightweight stripper, not a TeX parser: it removes
    comments and control sequences while keeping the human-readable argument
    text (section titles, bold/italic content, list items, hrefs), which is all
    the skills matcher needs.
    """
    text = data.decode("utf-8", errors="replace")
    text = _TEX_COMMENT_RE.sub("", text)
    # A TeX line break `\\` becomes a space before commands are stripped, so it
    # is not mistaken for an escaped character or a control word.
    text = text.replace("\\\\", " ")
    text = _TEX_HREF_RE.sub(r"\1", text)
    text = _TEX_URL_RE.sub(r"\1", text)
    text = _TEX_ENV_RE.sub(" ", text)
    # Repeat so one level of nesting (\textbf{\emph{X}}) unwraps to its content.
    for _ in range(3):
        new_text = _TEX_CMD_ARG_RE.sub(r" \1 ", text)
        if new_text == text:
            break
        text = new_text
    text = _TEX_CMD_BARE_RE.sub(" ", text)
    text = _TEX_ESCAPED_RE.sub(r"\1", text)
    text = text.replace("{", " ").replace("}", " ").replace("~", " ")
    return normalize_whitespace(text)
