"""Resume text extraction: PDF, DOCX, and .tex, plus the corrupt-file rule
that extraction must never raise (a bad upload yields "", never a 500)."""

import io

from app.extraction import extract_text, normalize_whitespace


def make_pdf(text: str) -> bytes:
    """A minimal, valid single-page PDF with one line of extractable text.

    Offsets and the xref table are computed here so pypdf parses it without
    falling back to reconstruction, keeping the test deterministic.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
    ]
    stream = b"BT /F1 24 Tf 72 700 Td (" + text.encode("latin-1") + b") Tj ET"
    objects.append(
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += str(index).encode() + b" 0 obj\n" + body + b"\nendobj\n"
    xref_pos = len(pdf)
    count = len(objects) + 1
    pdf += b"xref\n0 " + str(count).encode() + b"\n0000000000 65535 f \n"
    for offset in offsets:
        pdf += f"{offset:010d} 00000 n \n".encode()
    pdf += b"trailer\n<< /Size " + str(count).encode() + b" /Root 1 0 R >>\n"
    pdf += b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    return bytes(pdf)


def make_docx(paragraphs: list[str]) -> bytes:
    from docx import Document

    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_extract_pdf_returns_the_text():
    pdf = make_pdf("Python JavaScript Kubernetes Docker")
    assert extract_text("pdf", pdf) == "Python JavaScript Kubernetes Docker"


def test_extract_pdf_corrupt_returns_empty_not_raise():
    # Starts with the PDF magic but is otherwise garbage.
    assert extract_text("pdf", b"%PDF-1.4 this is not a real pdf body") == ""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def test_extract_docx_joins_paragraphs():
    docx = make_docx(["Python and Django experience", "AWS, Docker, Kubernetes"])
    text = extract_text("docx", docx)
    assert text == "Python and Django experience AWS, Docker, Kubernetes"


def test_extract_docx_reads_table_cells():
    from docx import Document

    document = Document()
    document.add_paragraph("Skills")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "PostgreSQL"
    table.rows[0].cells[1].text = "Redis"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("docx", buffer.getvalue())
    assert "PostgreSQL" in text
    assert "Redis" in text


def test_extract_docx_corrupt_returns_empty_not_raise():
    # PK zip magic but not a real docx package.
    assert extract_text("docx", b"PK\x03\x04 not actually a zip file") == ""


# ---------------------------------------------------------------------------
# TeX
# ---------------------------------------------------------------------------


def test_extract_tex_strips_commands_and_comments():
    tex = (
        b"% a hidden Fortran comment that must be ignored\n"
        b"\\documentclass{article}\n"
        b"\\begin{document}\n"
        b"\\section{Skills}\n"
        b"\\textbf{Python}, \\emph{Go}, and \\href{http://x}{PostgreSQL}.\n"
        b"\\begin{itemize}\n"
        b"\\item Kubernetes \\& Docker\n"
        b"\\end{itemize}\n"
        b"\\end{document}\n"
    )
    text = extract_text("tex", tex)
    # The section title and the formatted content survive as readable words.
    assert "Skills" in text
    assert "Python" in text
    assert "Go" in text
    assert "PostgreSQL" in text
    assert "Kubernetes" in text
    assert "Docker" in text
    # Control words and the comment are gone.
    assert "\\section" not in text
    assert "documentclass" not in text
    assert "Fortran" not in text
    assert "\\" not in text


def test_extract_tex_keeps_escaped_specials():
    tex = b"C\\# and F\\# plus 100\\% coverage"
    text = extract_text("tex", tex)
    assert "C#" in text
    assert "F#" in text
    assert "100% coverage" in text


def test_extract_tex_bad_bytes_do_not_raise():
    # Invalid UTF-8 is decoded with replacement, never an exception.
    assert isinstance(extract_text("tex", b"\xff\xfe\x00 Python"), str)


# ---------------------------------------------------------------------------
# General contract
# ---------------------------------------------------------------------------


def test_unknown_kind_and_empty_data_return_empty():
    assert extract_text("rtf", b"anything") == ""
    assert extract_text("pdf", b"") == ""


def test_normalize_whitespace_collapses_runs():
    assert normalize_whitespace("a  \n\t b   c ") == "a b c"
